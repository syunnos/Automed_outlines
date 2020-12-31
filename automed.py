import tkinter as tk

#---function--------------
import requests
from bs4 import BeautifulSoup
from docx import Document
import datetime

def ok_click():
    p_url = "https://pubmed.ncbi.nlm.nih.gov"
    s = text.get()
    s1= p_url+"/?term="+s
    soup = BeautifulSoup(requests.get(s1).content, "lxml")
    urls =[]
    atag = soup.select('a.docsum-title')
    for url in atag:
        urls.append(url['href'])

    #取得したインデックスをホームのURLに追加してそれぞれのページへ
    #wordファイルの生成
    word_write_file_name = s+str(datetime.date.today()) +".docx"
    doc = Document()

    class page_info:
        title =None
        authors =None
        date=None
        abstract =None
        journal =None

    for url in urls:
        paper = page_info()
        page = BeautifulSoup(requests.get(p_url+url).content, "lxml")
        paper.title =page.find('h1').text

        paper.abstract = page.find('div',id ="enc-abstract").text

        para = doc.add_paragraph()
        para.text =paper.title.strip()
        para = doc.add_paragraph()
        para.text = p_url + url
        para = doc.add_paragraph()
        para = doc.add_paragraph()
        para.text = paper.date

        para.text= paper.abstract.strip()
        doc.add_page_break()

    doc.save("C:/Users/syunnosukesuwa/Desktop/研究室/"+word_write_file_name)

#---GUI-------------------
win = tk.Tk()
win.title("Automed")
win.geometry("300x200")

label = tk.Label(win, text='Key words')
label.pack(anchor='center')

# テキストボックスを作成
text = tk.Entry(win)
text.pack(anchor='center')
text.insert(tk.END, '')
text.focus_set()

okButton = tk.Button(win, text='Get', command=ok_click)
okButton.pack(anchor='center',pady=5)

label = tk.Label(win, text='Use "+" in case of 2 or more words.')
label.pack(anchor='center')

win.mainloop()